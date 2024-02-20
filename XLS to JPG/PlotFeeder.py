import os
import datetime
import pandas as pd
import matplotlib.pyplot as plt
from docx.shared import Inches
import re
from docx import Document
from docx.shared import Pt
def parse_and_correct_datetime(date_str):
    try:
        parts = date_str.split(' ')
        date_str = ' '.join(parts[:-1])
        time_str = parts[-1]
        
        # Split the time string into hours and minutes
        hours, minutes = map(int, time_str.split(':'))
        
        # Correct minutes if they are 60
        if minutes == 60:
            minutes = 0
            hours += 1
        # Reconstruct the corrected datetime string
        corrected_dt_str = f'{date_str} {hours:02}:{minutes:02}'
        
        return corrected_dt_str
    except:
        return date_str
def sanitize_filename(filename):
    # Define a regular expression pattern to match reserved characters
    reserved_chars_pattern = r'[,\/:*?"<>|]'

    # Replace reserved characters with underscores
    sanitized_filename = re.sub(reserved_chars_pattern, '_', filename)

    return sanitized_filename

# Specify the directory containing the .xlsx files
directory_path = os.getcwd()
excel_file_pass=os.getcwd()

# Get a list of all .xlsx files in the directory
xlsx_files = [file for file in os.listdir(directory_path) if file.endswith('.xlsx')]
print(xlsx_files)
# Process each .xlsx file
i=0
ii=0
doc = Document()
HD1='Appendix 2: Load Data Provided By Redstone Arsenal'
doc.add_heading(HD1, level=1)
for file in xlsx_files:
    file_path = os.path.join(directory_path, file)
    print([i,'file is read'])
    print(['file name is',file])
    filen=file.replace('.xlsx',"")
    doc.add_heading(filen, level=2)
    filen=file.replace('.xlsx',"")
    print(['filen name is',filen])
    i=i+1
    # Read the Excel file into a Pandas DataFrame
    df = pd.read_excel(file_path)
    date_time_column = 'Time'
    # Apply the correction function to the DataFrame's column with date-time values
    df[date_time_column] = pd.to_datetime(df[date_time_column])
    #print(df[date_time_column])
    print(['The length of dates are:',len(df[date_time_column])])
    for des in range(1, len(df.columns)-1):
        print(des)
        print(['The length of data are:',len(df.iloc[:, des])])
        ########################## Smooth the Data ####################################
        window_size=5
        smoothed_data = df.iloc[:, des].rolling(window=window_size, center=True).median()
        max_value = smoothed_data.max()
        max_index = smoothed_data.idxmax()
        print('max index is:',max_index)
        print('max value is:',max_value)
        plt.figure(figsize=(10, 6))
        plt.plot(df[date_time_column], df.iloc[:, des], marker='o', linestyle='-', color='b')
        plt.annotate(f'Max: ({max_value})',
             xy=(df[date_time_column][max_index], max_value),
             xytext=(df[date_time_column][max_index], 3*max_value/4),
             arrowprops=dict(arrowstyle='->', color='red', linewidth=.5, shrinkA=1, shrinkB=1),
             fontsize=12, color='red', weight='bold')
        plt.title(['Plot of' ,df.columns[des], 'against Date and Time'])
        print(['Plot of' ,df.columns[des], 'against Date and Time'])
        name=sanitize_filename(f"{filen}{df.columns[des]}")
        print(name)
        plot_path = os.path.join(excel_file_pass, f"{name}{df.columns[des]}{'.jpg'}")    
        plt.xlabel('Date and Time')
        plt.ylabel('MW')
        plt.grid(True)
        print(plot_path)
        # Save the plot to the specified directory and filename
        plt.savefig(plot_path,dpi=600,bbox_inches='tight')
        doc.add_picture(plot_path, width=Inches(6))  # Adjust the width as needed
        table_caption = f"Figure {ii}: Plot of {df.columns[des]} against Date and Time"
        ii=ii+1
        caption_paragraph = doc.add_paragraph(table_caption)
        caption_paragraph.alignment = 1  # Center alignment
current_datetime = datetime.datetime.now()
timestamp = current_datetime.strftime("%Y-%m-%d_%H-%M-%S")
document_filename = output_file_path = f"{excel_file_pass}/Load Figures {timestamp}.docx"
doc.save(document_filename)